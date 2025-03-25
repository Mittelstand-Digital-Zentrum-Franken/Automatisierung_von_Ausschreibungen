# load data
import fitz
import json
import os
from PIL import Image
import pytesseract
import io
from typing import Tuple, List, Dict, Union
import base64
from io import BytesIO
import re

# Ollama
import ollama
ollama.pull(model='llama3.1')

# template
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

pytesseract.pytesseract.tesseract_cmd = r'C:\Users\sandra.nuissl\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'



def extract_text_from_pdf(pdf_document: fitz) -> list[str]:
    """Extracts the text from a PDF document and returns it as a list of texts and as context.

    Args:
        pdf_document: The PDF document from which the text is to be extracted. 

    Returns:
        context: A list containing a single string where all the text from the PDF pages is joined 
                together into a single context string.                         
    """

    # empty list
    texts = []

    # iterierate through the pages
    for page_num in range(len(pdf_document)):
        # load the current page
        page = pdf_document.load_page(page_num)

        # extract the text and append to the list (+ replace paragraph signs)
        text = page.get_text("text")
        text = text.replace("\n", "")
        texts.append(text)

    # join the texts to one context list
    context = [" ".join(texts)]

    return context


def extract_images_from_pdf(pdf_document: fitz) -> Tuple[str, Union[Image.Image, str]]:
    """ Extracts images from a PDF document and performs OCR to extract text from those images.

    Args:
        pdf_document: The PDF document from which the text is to be extracted. 

    Returns:
        images: A list of images extracted from the PDF.
        ocr_texts: A list of text strings extracted from the images.
    """

    # empty lists
    images = []
    ocr_texts = []

    # iterierate through the pages
    for page_num in range(len(pdf_document)):
        # load the current page
        page = pdf_document.load_page(page_num)

        # extract images from PDF
        image_list_per_page = page.get_images(full=True)
        for img_index, img in enumerate(image_list_per_page):
            xref = img[0]                                       # reference number
            base_image = pdf_document.extract_image(xref)       # extract image by reference number 
            image_bytes = base_image["image"]                   # extract image content

            # save image at the list as PIL
            image = Image.open(io.BytesIO(image_bytes))
            images.append(image)
            
            # extract content in image with OCR
            if image_bytes:
                try:
                    ocr_text = pytesseract.image_to_string(image)
                    ocr_texts.append(ocr_text)
                except Exception as e:
                    print(f"Error opening the image on page {page_num + 1}, Image {img_index + 1}: {e}")
            else:
                pass

    return images, ocr_texts


def process_pdf(pdf_path: str) -> Dict[str, Union[Image.Image, str]]:
    """Processes a PDF document to extract text, context, images, and OCR text from images.

    Args:
        pdf_path: The file path to the PDF document to be processed.

    Returns: 
        document_data: A dictionary containing the 'context', 'images' and 'ocr_texts'
    """

    # open PDF file
    pdf_document = fitz.open(pdf_path)

    # extract text
    context = extract_text_from_pdf(pdf_document)

    # extract images
    images, ocr_texts = extract_images_from_pdf(pdf_document)

    # close PDF file
    pdf_document.close()

    # store results in a dictionary
    document_data = {
        "context": context,
        "images": images,
        "ocr_texts": ocr_texts
    }

    return document_data


def get_summary(document_data: Dict[str, Union[Image.Image, str]]) -> Dict[str, List[str]]:
    """Create the summary of the uploaded pdf file.

    Args:
        document_data: A dictionary containing the 'context', 'images' and 'ocr_texts'

    Returns:
        bulletpoints_dict: A dictionary containing the extracted features of the categories depends on the extracted furniture
    """
        
    # Load example json
    example_json_path = './Data/example.json'
    with open(example_json_path, 'r', encoding='utf-8') as file:
        examples = json.load(file)
    print("examples load")

    # Get context
    text_to_summarize = document_data["context"]
    print("save context")

    # Get furniture
    modelfile_furniture = f'''
        Nenne den Möbeltyp, welcher im Text beschrieben wird in 2 Wörtern {text_to_summarize}?
        Verwende keine Produktnamen oder Firmennamen.
        Antworte mit nur 1 Subjektiv und nur einem erläuternden Adjektiv ohne Artikel
        Beispiele richtig: "Ergonomischer Schreibtischstuhl", "Höhenverstellbarer Schreibtisch", "Flexibler Besprechungstisch", ...
        Beispiele falsch: "Migration SE", "Steelcase", ...
        Antwortlänge = 2 Wörter
        '''

    furniture = ollama.generate(model='llama3.2', prompt=modelfile_furniture).response
    print("furniture finish")
    print(furniture)

    # Get criteria
    modelfile_criteria = f'''
        SYSTEM Deine Aufgabe ist es, Schlagwörter für die Kriterien, welche in einem technischen Vorbeschrieb für die Ausschreibung von einem vorgegebenen Möbelstück berücksichtigt werden müssen, aufzulisten\
        Nenne alle Kriterien, welche für das Möbelstück von bedeutung sind\
        Die Liste soll NUR die Schlagwörter enthalten, die als strukturierende Kategorien dienen könnten, ohne weitere Details oder erklärende Texte.\
        Hinweise:\
        - Keine numerischen Aufzählungen\
        - Nur Oberpunkte und Keine Unterpunkte\
        - Keine Anführungszeichen, wie "" oder ''\
        - keine Doppelpunkte, Klammern oder Spiegelstriche\
        - Keine Absätze\
        - keine erläuternden Sätze\
        - Keine einleitenden Worte wie "Hier ist die Liste mit den Schlagwörtern:"\
        - Keine abschließenden Worte wie "hier sind die Informationen ..." oder "ich kann dir noch weitere Informationen liefern ..."\
        - Mögliche Begriffe können sein: "Materialien", "Mechanik", "Nachhaltigkeit", "Zertifizierungen" usw.\
        - Der erste Punkt soll immer "Allgemein" sein, da dieser dann im nachgang befüllt wird\
        Erstelle eine Aufzählung aller Überschriften der technischen Kriterien für den technischen Vorbeschrieb welche für die Beschreibunge dieses Möbelstücks nötig sind: {furniture}
        '''

    criteria = ollama.generate(model='llama3.2', prompt=modelfile_criteria).response
    criteria_list = [re.sub(r"[^a-zA-ZäöüÄÖÜß\s-]", "", line).strip() for line in criteria.split("\n")]
    print("got criteria")

    # Get summary
    messages = f'''
        SYSTEM Your task is to create texts for the technical pre-description of tender documents for furniture based on the provided PDF document.\
        This is not a product advertisement but a neutral and technical presentation of the requirements.\
        Similar to these sample texts:\
        {examples}\
        Important: Do not mention company names or product names.\
        Focus exclusively on listing the contents of the document.\
        Write the text in the "should" perseptive and not in the "is" perspective\
        Notes:\
        - No numerical bullet points\
        - No escape sequences for tabs like t+\
        - No quotation marks\
        - No colons, parentheses, or dashes\
        - No paragraphs\
        - No formatting (e.g., bold text marked with **)\
        - No introductory phrases like "Here are the extracted contents of the text sections on the topic ..."\
        - No concluding phrases like "It should be noted ..." or "This information contains technical specifications ..."
        The provided text is {text_to_summarize}.\
        Extract from the provided text all content from the sections that relate to the topic {criteria_list}.\
        Exclude irrelevant sections. Focus on information that is directly or indirectly related to {criteria_list},\
        and present it in a technically precise manner.\
        Important: Do not mention company names or product names.
        '''

    summary = ollama.generate(model='llama3.2', prompt=messages)
    bulletpoints = summary['response'].split("\n")
    print("bulletpoints finish")

    # Save in dictionary
    bulletpoints_dict = {}
    current_key = None

    for line in bulletpoints:
        line = line.strip()
        # Skip empty line
        if not line:
            continue
        
        # Check if there is a key
        if line.startswith("**") and line.endswith("**"):
            current_key = re.sub(r"[*]", "", line).strip()
            bulletpoints_dict[current_key] = []
        elif current_key:
            line = re.sub(r"[*]", "", line).strip()
            bulletpoints_dict[current_key].append(line)

    return furniture, bulletpoints_dict


def create_document(massnahmennummer: str,
                    vergabenummer: str,
                    massnahme: str,
                    leistung: str,
                    produkt: str,
                    positionsnummer: str,
                    introduction: str,
                    hint: str,
                    #image: str,
                    output_path: str,
                    bulletpoints_dict: Dict[str, List[str]]) -> None:
    """Create a word document whith connect all data for the technical description.

    Args:
        massnahmennummer: Key Number 1
        vergabenummer: Key Number 2
        massnahme: Name of the projct
        leistung: Details to the project
        produkt: Product name
        positionsnummer: Position number
        introduction: Introduction
        hint: Additional information
        image: Image of the furniture
        output_path: Download name
        bulletpoints_dict: summary of the uploaded pdf file

    Returns:
        doc_stream: finish document as stream
    """
    # Create a new document
    doc = Document()

    # Adjust page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(60)
        section.bottom_margin = Pt(40)
        section.left_margin = Pt(60)
        section.right_margin = Pt(80)

    # Style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Create the header
    header = doc.sections[0].header
    paragraph_header = header.paragraphs[0]
    paragraph_header.text = (f"Maßnahmenummer: {massnahmennummer}\t \t Vergabenummer: {vergabenummer} \n"
                             f"Maßnahme: {massnahme} \nLeistung: {leistung}")
    paragraph_header.style = doc.styles["Header"]
    paragraph_header.style.font.size = Pt(9)

    # Add headline
    paragraph_headline = doc.add_heading(level=1)
    run_headline = paragraph_headline.add_run(f'Technischer Vorbeschrieb {produkt} (Pos. {positionsnummer})\n')
    run_headline.font.name = 'Arial'
    run_headline.font.size = Pt(12)
    run_headline.font.bold = True
    run_headline.font.color.rgb = RGBColor(0, 0, 0)
    run_headline.font.underline = True

    # First paragraph: introduction
    doc.add_paragraph(introduction)

    # Second paragraph: hint
    doc.add_paragraph(f'Hinweis zu Maßangaben: \n{hint} \n')

    # Third paragraph: reference
    paragraph_third = doc.add_paragraph()
    paragraph_third.add_run(f'Anforderungen an: {produkt}').bold = True
    doc.add_paragraph(
        'Das Produkt soll mindestens die folgenden Kriterien erfüllen: \n\n'
        'Alle Systemelemente sollen einer Designlinie entstammen. Das System muss serienmäßig lieferbar sein.')

    # Fourth paragraph: requirements
    paragraph_fourth = doc.add_paragraph()
    paragraph_fourth.add_run('\nAnforderung/ Kriterium').bold = True

    # List of requirements
    for category, bulletpoint in bulletpoints_dict.items():
        doc.add_paragraph(category, style='Normal')

        for point in bulletpoint:
            doc.add_paragraph(point, style='List Bullet')
        doc.add_paragraph("")

    # Image
    doc.add_paragraph('Beispielhafte Darstellung:')
    #doc.add_picture(image, width=Inches(4))

    # Add footer with page numbers
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Text
        run = paragraph.add_run("Seite ")
        run.font.size = Pt(10)

        # Current page
        fldSimple_current = OxmlElement('w:fldSimple')
        fldSimple_current.set(qn('w:instr'), 'PAGE')
        run._r.append(fldSimple_current)

        # Text
        run = paragraph.add_run(" von ")
        run.font.size = Pt(10)

        # Number for all pages
        fldSimple_total = OxmlElement('w:fldSimple')
        fldSimple_total.set(qn('w:instr'), 'NUMPAGES')
        run._r.append(fldSimple_total)

    # Save the document
    #doc.save(output_path)

    # Sent data to stream
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0) 

    return doc_stream
